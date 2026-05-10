"""
Create prepaidtemplate_tpl.docx by injecting Jinja2 placeholders
into the EXAMPLE block (col 0) of prepaidtemplate.docx.
Based on the EXACT 12-paragraph structure shown in the filled example.

Structure (col 0, 12 paras):
  Para 0  sz=28: "To:"                               -- untouched
  Para 1  sz=28: Name                                -- {{ bN_name }}
  Para 2  sz=28: Full address (Word wraps it)        -- {{ bN_addr }}
  Para 3  sz=28: "Pin:XXXXXX,"                       -- {{ bN_pin }}
  Para 4  sz=28: " Mob:XXXXXXXXXX"                   -- {{ bN_mob }}
  Para 5  sz=28: ""  (gap)                           -- untouched
  Para 6  sz=28: "[spaces]From:"                     -- untouched
  Para 7  sz=24: "[items][spaces]CREAM X EMIRATES "  -- {{ bN_order }} + spaces + CREAM
  Para 8  sz=24: "[spaces]PUTHUPALLY, KTM"           -- untouched
  Para 9  sz=24: "[spaces]Pin: 686011"               -- untouched
  Para 10 sz=24: "[spaces]Mob: 8129770502"           -- untouched
  Para 11 sz=24: "Biller ID: XXXXXXXXXX"             -- {{ bN_biller }}

The blank blocks (col 1, other rows) use a 15-para structure — we copy the
col-0 structure to all slots so every block is consistent.
"""

import copy
from docx import Document
from docx.oxml.ns import qn

def get_text(para):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    return ''.join(t.text or '' for t in para.findall('.//w:t', ns))

def set_run_text(para, new_text):
    """Replace the <w:t> content of the first run, keep ALL formatting."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    runs = para.findall(qn('w:r'))
    if runs:
        t_el = runs[0].find(qn('w:t'))
        if t_el is None:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            t_el = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve"></w:t>')
            runs[0].append(t_el)
        t_el.text = new_text
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # Remove extra runs
        for r in runs[1:]:
            para.remove(r)
    else:
        # No run: copy from para 1 of example (same formatting)
        ref = example_paras[1].findall(qn('w:r'))
        if ref:
            new_r = copy.deepcopy(ref[0])
            t_el = new_r.find(qn('w:t'))
            if t_el is not None:
                t_el.text = new_text
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            para.append(new_r)


doc = Document('prepaidtemplate.docx')
table = doc.tables[0]
rows = table._tbl.findall(qn('w:tr'))

# Read the example block (row 0, col 0) as reference
example_cell = rows[0].findall(qn('w:tc'))[0]
example_paras = example_cell.findall(qn('w:p'))

# Para 7 original: " 1 CXE                                                   CREAM X EMIRATES "
# Extract the CREAM X EMIRATES part (right side after items)
p7_text = get_text(example_paras[7])
# Find where CREAM starts (after the leading spaces+items)
cream_idx = p7_text.find('CREAM X EMIRATES')
# Leading spaces before "CREAM" section — use fixed spacing from original
# The pattern: {items}{spaces}CREAM X EMIRATES
# We'll use a fixed number of spaces between order and CREAM
ORDER_TO_CREAM_SPACES = p7_text[p7_text.find(' '):cream_idx]  # spaces between end of items and CREAM

# All block positions: (row_index, col_index)
block_positions = [
    (0, 0), (0, 1),
    (1, 0), (1, 1),
    (2, 0), (2, 1),
]

for block_idx, (row_i, col_i) in enumerate(block_positions):
    b = f"b{block_idx}"
    row = rows[row_i]
    cell = row.findall(qn('w:tc'))[col_i]

    # Replace cell contents with a deep copy of the example block
    # Remove all existing paragraphs
    existing_paras = cell.findall(qn('w:p'))
    for p in existing_paras:
        cell.remove(p)

    # Copy all 12 paragraphs from example block
    for src_para in example_paras:
        new_para = copy.deepcopy(src_para)
        cell.append(new_para)

    # Now set placeholders in the copied paragraphs
    cell_paras = cell.findall(qn('w:p'))

    # Para 1: Name
    set_run_text(cell_paras[1], f"{{{{ {b}_name }}}}")

    # Para 2: Address (Word wraps automatically)
    set_run_text(cell_paras[2], f"{{{{ {b}_addr }}}}")

    # Para 3: Pin line
    set_run_text(cell_paras[3], f"{{{{ {b}_pin }}}}")

    # Para 4: Mob line
    set_run_text(cell_paras[4], f"{{{{ {b}_mob }}}}")

    # Para 5: keep empty gap — untouched

    # Para 6: keep "[spaces]From:" — untouched

    # Para 7: "{{ bN_order }}[spaces]CREAM X EMIRATES "
    # Rebuild: placeholder for order on left, fixed CREAM text on right
    cream_right = "CREAM X EMIRATES "
    # Use same spacing ratio as original — 53 spaces between order placeholder end and CREAM
    # We'll use a fixed large space block that mimics the original
    spaces_between = " " * 50  # approximate, will be close enough
    set_run_text(cell_paras[7], f"{{{{ {b}_order }}}}{spaces_between}{cream_right}")

    # Para 8-10: keep PUTHUPALLY, Pin, Mob — untouched

    # Para 11: Biller ID
    set_run_text(cell_paras[11], f"{{{{ {b}_biller }}}}")

doc.save('prepaidtemplate_tpl.docx')
print("Created prepaidtemplate_tpl.docx")

# Verify
doc2 = Document('prepaidtemplate_tpl.docx')
table2 = doc2.tables[0]
rows2 = table2._tbl.findall(qn('w:tr'))
cell2 = rows2[0].findall(qn('w:tc'))[0]
paras2 = cell2.findall(qn('w:p'))
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
print("Block 0 placeholders:")
for i, p in enumerate(paras2):
    t = ''.join(x.text or '' for x in p.findall('.//w:t', ns))
    print(f"  Para {i:2d}: [{t[:80]}]")
