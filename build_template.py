"""
build_template.py
=================
Builds 'prepaidtemplate_tpl.docx' from 'kunjii.docx'.

Matches the EXACT template layout shown by the user:

  To:
  {{ bN_name }}
  {{ bN_addr }}         <- street address (may wrap)
  {{ bN_pin }}          <- "Pin: 600094"
  {{ bN_state }}        <- "Tamil Nadu"
  {{ bN_mob }}          <- "Mob: 9790551399"
  [blank]
  {{ bN_order }}        <- "3CXE" / "1 CXE, 1 BL"
  [blank]
  [spaces]From:
  [spaces]CREAM X EMIRATES
  [spaces]PUTHUPALLY, KTM
  [spaces]Pin: 686011
  [spaces]Mob: 8129770502
  Biller ID: {{ bN_biller }}

Sources:
  - Paras 0-4 cloned from Cell[0][0] (correct addr formatting)
  - From section cloned from Cell[0][1] (Paras 7-12, correct 73-char spacing)
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

TEMPLATE_IN  = 'kunjii.docx'
TEMPLATE_OUT = 'prepaidtemplate_tpl.docx'


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def get_text(para_el):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    return ''.join(t.text or '' for t in para_el.findall('.//w:t', ns))


def first_run(para_el):
    """Return the first <w:r> in a paragraph element, or None."""
    runs = para_el.findall(qn('w:r'))
    return runs[0] if runs else None


def set_text(para_el, new_text, clone_run_from=None):
    """
    Set the text content of a paragraph's first run, preserving all formatting.
    If the paragraph has no runs, clone one from clone_run_from.
    """
    runs = para_el.findall(qn('w:r'))

    if not runs:
        if clone_run_from is not None:
            new_run = copy.deepcopy(clone_run_from)
            t_ex = new_run.find(qn('w:t'))
            if t_ex is not None:
                new_run.remove(t_ex)
        else:
            new_run = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        para_el.append(new_run)
        runs = [new_run]

    # Keep first run, drop the rest
    keep = runs[0]
    for r in runs[1:]:
        para_el.remove(r)

    t_el = keep.find(qn('w:t'))
    if t_el is None:
        t_el = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve"></w:t>')
        keep.append(t_el)

    t_el.text = new_text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def blank_para_like(ref_para_el):
    """Return a deep-copy of ref_para_el with its run text cleared."""
    p = copy.deepcopy(ref_para_el)
    for t in p.findall('.//' + qn('w:t')):
        t.text = ''
    return p


# ---------------------------------------------------------------------------
# Cell builder
# ---------------------------------------------------------------------------

def build_cell(cell_xml, addr_paras, from_paras, b):
    """
    Populate one label cell with the correct 15-paragraph structure.

    addr_paras  : Cell[0][0] Paras 0-4  (To:, name, addr+pin, state+mob, blank)
    from_paras  : Cell[0][1] Paras 7-12 (From section with correct spacing)
    b           : block prefix e.g. 'b0'
    """
    # Clear existing paragraphs
    for p in cell_xml.findall(qn('w:p')):
        cell_xml.remove(p)

    # ── Step 1: add address section from Cell[0][0] Paras 0-4 ───────────────
    #   [0] To:
    #   [1] {Name}       → placeholder
    #   [2] {Addr+Pin}   → addr placeholder  (we strip old placeholder text)
    #   [3] {State+Mob}  → mob placeholder   (we strip old placeholder text)
    #   [4] blank
    for src in addr_paras:
        cell_xml.append(copy.deepcopy(src))

    paras = cell_xml.findall(qn('w:p'))
    # Reference run for cloning (use name para's run — carries correct font)
    ref_run = first_run(paras[1])

    # Set placeholders for existing paras
    set_text(paras[1], f'{{{{ {b}_name }}}}')          # Para 1 → name
    set_text(paras[2], f'{{{{ {b}_addr }}}}')           # Para 2 → street address
    set_text(paras[3], f'{{{{ {b}_mob }}}}', ref_run)   # Para 3 → mobile (repurposed)
    # Para 4 stays blank

    # ── Step 2: insert Pin and State lines after Para 2 ─────────────────────
    # Clone Para 2 twice (preserves formatting) and insert before Para 3
    pin_para   = copy.deepcopy(paras[2])
    state_para = copy.deepcopy(paras[2])
    paras[3].addprevious(pin_para)    # → becomes Para 3
    paras[3].addprevious(state_para)  # → becomes Para 4; original Para 3 shifts to Para 5

    # Refresh after insertions
    paras = cell_xml.findall(qn('w:p'))
    # Now: [0]To: [1]name [2]addr [3]pin [4]state [5]mob [6]blank
    set_text(paras[3], f'{{{{ {b}_pin }}}}',   ref_run)  # Para 3 → Pin:600094
    set_text(paras[4], f'{{{{ {b}_state }}}}', ref_run)  # Para 4 → State

    # ── Step 3: insert Order line and blank after Para 6 (blank) ────────────
    # Para 6 is the blank gap from addr_paras[4]
    order_para = copy.deepcopy(paras[6])   # clone blank → becomes order
    gap_para   = copy.deepcopy(paras[6])   # clone blank → gap after order

    # Insert both AFTER Para 6:  ...[6]blank [7]order [8]gap [9]From:...
    # To insert after Para 6, we add before Para 6's next sibling.
    # Since Para 6 is last right now, just append them to cell_xml before From section.
    cell_xml.append(order_para)
    cell_xml.append(gap_para)

    # Refresh and set order
    paras = cell_xml.findall(qn('w:p'))
    # [0]To [1]name [2]addr [3]pin [4]state [5]mob [6]blank [7]order [8]gap
    set_text(paras[7], f'{{{{ {b}_order }}}}', ref_run)  # Para 7 → 3CXE

    # ── Step 4: append From section from Cell[0][1] Paras 7-12 ──────────────
    # [9]From: [10]CREAM X EMIRATES [11]PUTHUPALLY [12]Pin:686011
    # [13]Mob:8129770502 [14]Biller ID
    for src in from_paras:
        cell_xml.append(copy.deepcopy(src))

    # ── Step 5: set biller placeholder on last para (index 14) ──────────────
    paras = cell_xml.findall(qn('w:p'))
    set_text(paras[14], f'{{{{ {b}_biller }}}}')

    return len(paras)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document(TEMPLATE_IN)
    table = doc.tables[0]

    # Reference paragraphs
    cell00_xml = table.rows[0].cells[0]._tc
    cell01_xml = table.rows[0].cells[1]._tc

    all00 = cell00_xml.findall(qn('w:p'))
    all01 = cell01_xml.findall(qn('w:p'))

    addr_paras = all00[0:5]    # Cell[0][0] Paras 0-4
    from_paras = all01[7:13]   # Cell[0][1] Paras 7-12 (correct From spacing)

    print('Addr paras (Cell[0][0] 0-4):')
    for i, p in enumerate(addr_paras):
        print(f'  [{i}] {repr(get_text(p)[:60])}')
    print('From paras (Cell[0][1] 7-12):')
    for i, p in enumerate(from_paras, 7):
        print(f'  [{i}] {repr(get_text(p)[:70])}')

    block_positions = [(0,0),(0,1),(1,0),(1,1),(2,0),(2,1)]

    for block_idx, (ri, ci) in enumerate(block_positions):
        b = f'b{block_idx}'
        cell_xml = table.rows[ri].cells[ci]._tc
        n = build_cell(cell_xml, addr_paras, from_paras, b)
        print(f'  Block {block_idx} ({ri},{ci}) -> {n} paras')

    doc.save(TEMPLATE_OUT)
    print(f'\n[DONE] Saved: {TEMPLATE_OUT}')

    # Verify
    print('\n--- Verification ---')
    doc2 = Document(TEMPLATE_OUT)
    table2 = doc2.tables[0]
    labels = ['To:', 'name', 'addr', 'pin', 'state', 'mob',
              'blank', 'order', 'blank2',
              'From:', 'CREAM', 'PUTHUPALLY', 'Pin:686011', 'Mob:812',
              'biller']
    for block_idx, (ri, ci) in enumerate(block_positions):
        cell = table2.rows[ri].cells[ci]
        paras = cell.paragraphs
        print(f'\n  Block {block_idx} ({len(paras)} paras):')
        for i, (p, lbl) in enumerate(zip(paras, labels)):
            print(f'    [{i:2}] {lbl:<12} = [{p.text[:45]}]')


if __name__ == '__main__':
    main()
